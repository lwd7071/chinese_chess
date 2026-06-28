import base64
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "BAO_CAO_AI_source.md"
OUT = ROOT / "Bao_cao_AI_Co_Vua_6_Level.docx"
IMG_DIR = ROOT / "report_images"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(33, 37, 41)
MUTED = RGBColor(90, 90, 90)
TABLE_FILL = "F2F4F7"
FORMULA_FILL = "F7F9FC"


def clean_inline(text: str) -> str:
    text = re.sub(r"\{#[^}]+\}", "", text)
    text = text.replace("\\-", "-").replace("\\_", "_").replace("\\*", "*")
    text = text.replace("â€“", "-").replace("â†’", "->").replace("â€”", "-")
    text = text.replace("–", "-").replace("—", "-").replace("→", "->")
    text = text.replace("📷", "Hình:").replace("⚠️", "Lưu ý:")
    text = text.replace("ðŸ“·", "Hình:").replace("âš ï¸", "Lưu ý:")
    text = text.replace("✅", "Đạt").replace("❌", "Không đạt")
    text = text.replace("âœ…", "Đạt").replace("âŒ", "Không đạt")
    text = re.sub(r"[\U0001F000-\U0001FAFF]", "", text)
    text = text.replace("\u200b", "")
    return text.strip()


def strip_md_marks(text: str) -> str:
    text = clean_inline(text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = text.replace("\\", "")
    return text.strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_run(run, font="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_rich_text(paragraph, text: str, base_size=11, color=INK):
    text = clean_inline(text)
    if not text:
        return
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            style_run(run, size=base_size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            style_run(run, size=base_size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            style_run(run, font="Consolas", size=max(base_size - 1, 8), color=DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        style_run(run, size=base_size, color=color)


def add_paragraph(doc, text: str, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    add_rich_text(p, text)
    return p


def add_formula(doc, formula: str):
    formula = clean_inline(formula)
    if not formula:
        return
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9000, 180)
    set_cell_margins(table, top=100, bottom=100, start=180, end=180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, FORMULA_FILL)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula.replace("\\\\", "\\"))
    style_run(run, font="Cambria Math", size=11, color=DARK_BLUE)
    doc.add_paragraph()


def table_rows(lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if re.fullmatch(r"\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?", stripped):
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            parts = [strip_md_marks(c) for c in stripped.strip("|").split("|")]
            if len(parts) >= 2:
                rows.append(parts)
    if not rows:
        return []
    width = max(len(r) for r in rows)
    return [r + [""] * (width - len(r)) for r in rows]


def add_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_cell_margins(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, value, base_size=9 if cols >= 4 else 10)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = DARK_BLUE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif cols <= 3 and c_idx == 0 and len(value) < 24:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def extract_images(text: str):
    IMG_DIR.mkdir(exist_ok=True)
    images = {}
    pattern = re.compile(r"^\[(image\d+)\]:\s*<data:image/(\w+);base64,([^>]+)>", re.M)
    for match in pattern.finditer(text):
        name, ext, data = match.groups()
        path = IMG_DIR / f"{name}.{ext.lower()}"
        path.write_bytes(base64.b64decode(data))
        images[name] = path
    text = pattern.sub("", text)
    return text, images


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRƯỜNG ĐẠI HỌC SƯ PHẠM KỸ THUẬT TP. HCM\nKHOA CÔNG NGHỆ THÔNG TIN")
    style_run(r, size=13, color=DARK_BLUE, bold=True)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DỰ ÁN: CỜ VUA 6 LEVEL")
    style_run(r, size=24, color=BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Báo cáo cuối kỳ")
    style_run(r, size=18, color=INK, bold=True)

    doc.add_paragraph()
    for line in [
        "Môn học: Trí Tuệ Nhân Tạo",
        "Mã số lớp HP: ARIN",
        "GVHD: Phan Thị Huyền Trang",
        "Nhóm thực hiện: Nhóm 1",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        style_run(r, size=12, color=INK, bold=True)

    doc.add_paragraph()
    members = doc.add_paragraph()
    members.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = members.add_run("Thành viên\nTrần Lê Thái - 24110331\nLương Viết Vĩ Đông - 24110202\nNguyễn Minh Trí - 24110359")
    style_run(r, size=12, color=INK)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TP. HỒ CHÍ MINH - THÁNG 7 / NĂM 2026")
    style_run(r, size=12, color=MUTED, bold=True)
    doc.add_page_break()


def add_teacher_page(doc):
    doc.add_heading("NHẬN XÉT VÀ GHI ĐIỂM CỦA GIÁO VIÊN CHẤM", level=1)
    rows = [
        ["Nội dung", "Nhận xét"],
        ["Tính đúng đắn của chương trình", ""],
        ["Mức độ hiểu và chạy tay thuật toán", ""],
        ["Thực nghiệm, ảnh động, đánh giá", ""],
        ["Hình thức báo cáo", ""],
        ["Điểm số", ""],
        ["Chữ ký giáo viên", ""],
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_cell_margins(table, top=140, bottom=140, start=140, end=140)
    heights = [0.35, 0.75, 0.75, 0.75, 0.75, 0.5, 0.9]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_rich_text(p, value, base_size=10)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = DARK_BLUE
            tr_pr = cell._tc.getparent().get_or_add_trPr()
            tr_h = OxmlElement("w:trHeight")
            tr_h.set(qn("w:val"), str(int(heights[r_idx] * 1440)))
            tr_h.set(qn("w:hRule"), "atLeast")
            tr_pr.append(tr_h)
    doc.add_page_break()


def collect_toc(lines):
    entries = []
    for line in lines:
        s = line.strip()
        if not s.startswith("#"):
            continue
        level = len(s) - len(s.lstrip("#"))
        title = strip_md_marks(s.lstrip("#").strip())
        if title and len(title) < 120 and not title.startswith("self.") and "```" not in title:
            entries.append((min(level, 3), title))
    return entries


def add_toc(doc, entries):
    doc.add_heading("MỤC LỤC", level=1)
    seen = set()
    for level, title in entries:
        if title in seen:
            continue
        seen.add(title)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 * (level - 1))
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        style_run(r, size=10, color=INK)
    doc.add_page_break()


def add_body(doc, lines, images):
    i = 0
    in_formula = False
    formula = []
    used_images = set()
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("**TRƯỜNG") or line.startswith("**DỰ ÁN") or line.startswith("**Báo cáo") or line.startswith("**Môn học"):
            i += 1
            continue
        if line.startswith("**NHẬN XÉT") or line.startswith("| Nội dung | Nhận xét") or line.startswith("**MỤC LỤC") or re.match(r"^\[.*\]\(#", line):
            i += 1
            continue
        if re.match(r"^\[image\d+\]:", line):
            i += 1
            continue

        if "$$" in line:
            before, _, after = line.partition("$$")
            if before.strip():
                add_paragraph(doc, before)
            if in_formula:
                formula.append(after)
                add_formula(doc, " ".join(formula))
                formula = []
                in_formula = False
            else:
                in_formula = True
                formula = [after]
            i += 1
            continue
        if in_formula:
            formula.append(line)
            i += 1
            continue

        img_ref = re.fullmatch(r"!\[\]\[(image\d+)\]", line)
        if img_ref is None:
            img_ref = re.fullmatch(r"!\[[^\]]*\]\[(image\d+)\]", line)
        if img_ref:
            img = images.get(img_ref.group(1))
            if img and img.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img), width=Inches(5.7))
                used_images.add(img_ref.group(1))
            i += 1
            continue

        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, table_rows(block))
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = strip_md_marks(line.lstrip("#").strip())
            if title and len(title) < 120 and not title.startswith("self.") and "`" not in title:
                doc.add_heading(title, level=min(level, 3))
            else:
                add_paragraph(doc, title)
            i += 1
            continue

        if line.startswith(("* ", "- ")):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
            i += 1
            continue

        numbered = re.match(r"^\d+[.)]\s+(.*)", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, numbered.group(1))
            i += 1
            continue

        if line in {"---", "***"}:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("—")
            style_run(r, size=11, color=MUTED)
            i += 1
            continue

        add_paragraph(doc, line)
        i += 1
    missing = [name for name in sorted(images) if name not in used_images]
    if missing:
        doc.add_heading("Minh họa giao diện chương trình", level=2)
        for name in missing:
            img = images[name]
            if img.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img), width=Inches(5.7))


def main():
    text = SOURCE.read_text(encoding="utf-8")
    text = text.replace("\x0b", " ").replace("\r\n", "\n")
    text, images = extract_images(text)
    lines = text.splitlines()

    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_teacher_page(doc)
    add_toc(doc, collect_toc(lines))

    section = doc.sections[-1]
    header = section.header.paragraphs[0]
    header.text = "Báo cáo cuối kỳ - Trí Tuệ Nhân Tạo"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    add_body(doc, lines, images)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
